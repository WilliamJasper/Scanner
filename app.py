import io
import os
import re
import json
import asyncio
import datetime
import threading
import time
from pathlib import Path
from typing import Optional

from docx import Document

from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException, Depends
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import pdfplumber
import httpx
import pyodbc

# Database imports
from sqlalchemy import create_engine, Column, Integer, String, Float, DateTime, ForeignKey, Text, Unicode, UnicodeText
from sqlalchemy.orm import sessionmaker, Session, relationship, declarative_base

load_dotenv()

app = FastAPI(title="ScannerTTB - Bill For Collection Scanner")

# Setup templates
BASE_DIR = Path(__file__).resolve().parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

# Static files
static_dir = BASE_DIR / "static"
static_dir.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")

# Upload directory
upload_dir = BASE_DIR / "uploads"
upload_dir.mkdir(exist_ok=True)

TYPHOON_API_KEY = os.getenv("TYPHOON_API_KEY", "")
TYPHOON_API_URL = "https://api.opentyphoon.ai/v1/chat/completions"
TYPHOON_MODEL = "typhoon-v2.5-30b-a3b-instruct"
TYPHOON_OCR_URL = "https://api.opentyphoon.ai/v1/ocr"
TYPHOON_OCR_MODEL = os.getenv("TYPHOON_OCR_MODEL", "typhoon-ocr")
TAX_OCR_JOBS: dict[str, dict] = {}
TAX_OCR_JOBS_LOCK = threading.Lock()

# ======== DATABASE SETUP ========
# Connection string for SQL Server (Trusted Connection)
DATABASE_URL = os.getenv("DATABASE_URL", "mssql+pyodbc://LAPTOP-2CN8L0R4\\SQLEXPRESS/ScannerDB?driver=ODBC+Driver+17+for+SQL+Server&trusted_connection=yes")

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class Bill(Base):
    __tablename__ = "Bills"
    Id = Column(Integer, primary_key=True, index=True)
    BillNumber = Column("เลขที่บิล", Unicode(100))
    CustomerName = Column("ชื่อลูกค้า", Unicode(255))
    BillDate = Column("วันที่", Unicode(50))
    Address = Column("ที่อยู่", UnicodeText)
    CreditLineNo = Column("เลขวงเงิน", Unicode(100))
    Ref1 = Column("อ้างอิง1", Unicode(100))
    Ref2 = Column("อ้างอิง2", Unicode(100))
    TotalAmount = Column("ยอดเงินรวม", Float)
    CalculatedTotal = Column("ยอดเงินคำนวณ", Float)
    BankName = Column("ชื่อธนาคาร", Unicode(255))
    TotalPages = Column("จำนวนหน้าทั้งหมด", Integer)
    CreatedAt = Column("วันที่สร้าง", DateTime, default=datetime.datetime.utcnow)
    
    items = relationship("BillItem", back_populates="bill")
    verification = relationship("Verification", back_populates="bill", uselist=False)

class BillItem(Base):
    __tablename__ = "BillItems"
    Id = Column(Integer, primary_key=True, index=True)
    BillId = Column("ลำดับ", Integer, ForeignKey("Bills.Id"))
    ChassisNumber = Column("เลขตัวถัง", Unicode(100))
    Status = Column("สถานะ", Unicode(50))
    Description = Column("รายการ", Unicode(255))
    Period = Column("ระยะเวลา", Unicode(50))
    Days = Column("จำนวนวัน", Unicode(50))
    InterestRate = Column("อัตราดอกเบี้ย", Unicode(50))
    DueDate = Column("วันครบกำหนด", Unicode(50))
    Principal = Column("เงินต้น", Unicode(100))
    AmountDue = Column("ยอดที่ต้องชำระ", Float)
    PageNumber = Column("หน้าที่", Integer)
    
    bill = relationship("Bill", back_populates="items")

class Verification(Base):
    __tablename__ = "Verifications"
    Id = Column(Integer, primary_key=True, index=True)
    BillId = Column("ไอดีบิล", Integer, ForeignKey("Bills.Id"))
    DataCompleteness = Column("ความสมบูรณ์ของข้อมูล", Unicode(50))
    AmountCheck = Column("ตรวจสอบยอดเงิน", Unicode(255))
    Notes = Column("หมายเหตุ", UnicodeText)
    Recommendation = Column("ข้อเสนอแนะ", UnicodeText)
    
    bill = relationship("Bill", back_populates="verification")

# Create tables if they don't exist
Base.metadata.create_all(bind=engine)

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def extract_from_pdf(pdf_path: str, max_pages: int = 0, password: str = "") -> dict:
    """Extract structured data from Bill For Collection PDF using pdfplumber table extraction.

    Args:
        pdf_path: Path to the PDF file
        max_pages: Maximum number of pages to process (0 = all pages)
        password: Password for encrypted PDF (empty string if not encrypted)

    Returns:
        dict with bill_info, items, summary, and raw_text
    """
    bill_info = {
        "bill_number": "",
        "customer_name": "",
        "date": "",
        "address": "",
        "credit_line_no": "",
        "ref1": "",
        "ref2": "",
    }
    items = []
    raw_text_pages = []
    total_amount = ""
    total_amount_text = ""
    bank_name = ""
    total_pdf_pages = 0

    open_kwargs = {}
    if password:
        open_kwargs["password"] = password

    with pdfplumber.open(pdf_path, **open_kwargs) as pdf:
        total_pdf_pages = len(pdf.pages)
        pages_to_process = total_pdf_pages if max_pages <= 0 else min(max_pages, total_pdf_pages)

        for page_idx in range(pages_to_process):
            page = pdf.pages[page_idx]
            page_text = page.extract_text() or ""
            raw_text_pages.append(page_text)

            # --- Extract header info from first page ---
            if page_idx == 0:
                bill_info = _parse_header(page_text)

            # --- Extract table data ---
            table = page.extract_table()
            if table:
                for row_idx, row in enumerate(table):
                    # Skip header row (first row of each page table)
                    if row_idx == 0:
                        continue

                    # Ensure row has enough columns
                    if len(row) < 8:
                        continue

                    chassis_raw = (row[0] or "").strip()
                    period = (row[2] or "").strip()
                    days = (row[3] or "").strip()
                    interest_rate = (row[4] or "").strip()
                    due_date = (row[5] or "").strip()
                    principal = (row[6] or "").strip()
                    amount_due = (row[7] or "").strip()

                    # Skip if no chassis number
                    if not chassis_raw:
                        continue

                    # Parse chassis number and status
                    chassis_number, status = _parse_chassis_status(chassis_raw)

                    # Skip if chassis doesn't look valid
                    if not re.match(r"^[A-Z0-9]{10,}", chassis_number):
                        continue

                    items.append({
                        "chassis_number": chassis_number,
                        "status": status,
                        "description": "ดอกเบี้ย",
                        "period": period,
                        "days": days,
                        "interest_rate": interest_rate,
                        "due_date": due_date,
                        "principal": principal,
                        "amount_due": amount_due,
                        "page_number": page_idx + 1,
                    })

            # --- Extract total from last page ---
            if page_idx == pages_to_process - 1:
                total_amount, total_amount_text, bank_name = _parse_footer(page_text)

    # Calculate totals
    closed_count = sum(1 for item in items if item["status"] == "Closed")
    active_count = sum(1 for item in items if item["status"] == "Active")
    
    # Calculate sum from items
    calculated_total = 0.0
    for item in items:
        try:
            calculated_total += float(item["amount_due"].replace(",", ""))
        except (ValueError, AttributeError):
            pass

    raw_text = "\n\n--- หน้า ---\n\n".join(raw_text_pages)

    return {
        "bill_info": bill_info,
        "items": items,
        "summary": {
            "total_items": len(items),
            "total_amount": total_amount or f"{calculated_total:,.2f}",
            "total_amount_text": total_amount_text,
            "closed_items": closed_count,
            "active_items": active_count,
            "bank_name": bank_name,
            "calculated_total": f"{calculated_total:,.2f}",
            "total_pdf_pages": total_pdf_pages,
            "pages_processed": pages_to_process,
        },
        "raw_text": raw_text,
    }


def _parse_chassis_status(raw: str) -> tuple:
    """Parse chassis number and Closed/Active status from raw cell value."""
    raw = raw.replace("\n", " ").strip()
    if "(Closed" in raw:
        chassis = raw.split("(Closed")[0].strip()
        return chassis, "Closed"
    return raw.strip(), "Active"


def _parse_header(text: str) -> dict:
    """Parse header information from page text."""
    info = {
        "bill_number": "",
        "customer_name": "",
        "date": "",
        "address": "",
        "credit_line_no": "",
        "ref1": "",
        "ref2": "",
    }

    # Bill number
    m = re.search(r"Bill For Collection\s+.*?:\s*(\d+)", text)
    if m:
        info["bill_number"] = m.group(1)

    # Date
    m = re.search(r"(\d{1,2}/\d{1,2}/\d{4})\s*$", text, re.MULTILINE)
    if m:
        info["date"] = m.group(1)

    # Credit Line No
    m = re.search(r"Credit Line No\.\s*:\s*(\S+)", text)
    if m:
        info["credit_line_no"] = m.group(1)

    # Ref 1
    m = re.search(r"Ref 1\s*:\s*(\S+)", text)
    if m:
        info["ref1"] = m.group(1)

    # Ref 2
    m = re.search(r"Ref 2\s*:\s*(\S+)", text)
    if m:
        info["ref2"] = m.group(1)

    # For Thai text fields, use the text as-is from the PDF reader
    # (extract_text may have encoding issues in console but data is correct)
    lines = text.split("\n")
    for line in lines:
        # Customer name line contains the date at the end
        if ":" in line and re.search(r"\d{1,2}/\d{1,2}/\d{4}", line):
            # This might be the customer name line
            parts = line.split(":")
            if len(parts) >= 2 and "Bill" not in line and "Credit" not in line:
                name_part = parts[1].strip()
                # Remove date part
                name_part = re.sub(r"\s+.*\d{1,2}/\d{1,2}/\d{4}", "", name_part).strip()
                if name_part:
                    info["customer_name"] = name_part

        # Address line
        if ":" in line and "Credit Line" in line:
            addr_part = line.split(":")[1].strip()
            addr_part = re.sub(r"\s*Credit Line.*", "", addr_part).strip()
            if addr_part:
                info["address"] = addr_part

    return info


def _parse_footer(text: str) -> tuple:
    """Parse footer from last processed page to get total amount."""
    total_amount = ""
    total_amount_text = ""
    bank_name = ""

    # Total amount pattern: number at end after parenthesized text
    m = re.search(r"\((.+?)\)\s+([\d,]+\.\d{2})\s*$", text, re.MULTILINE)
    if m:
        total_amount_text = m.group(1)
        total_amount = m.group(2)

    # Bank name
    m = re.search(r"(ธนาคาร.+?)(?:\n|$)", text)
    if not m:
        # Try to find TTB reference
        if "TTB" in text or "ทหารไทย" in text or "ธนชาต" in text:
            bank_name = "ธนาคารทหารไทยธนชาต จำกัด (มหาชน)"
    else:
        bank_name = m.group(1).strip()

    return total_amount, total_amount_text, bank_name


async def verify_with_typhoon(parsed_data: dict) -> dict:
    """Send a compact summary to Typhoon AI for verification."""

    summary_text = f"""ข้อมูลที่แยกได้จากเอกสาร Bill For Collection ของธนาคาร TTB:

ข้อมูล Bill:
- เลขที่: {parsed_data['bill_info']['bill_number']}
- วันที่: {parsed_data['bill_info']['date']}
- Credit Line No: {parsed_data['bill_info']['credit_line_no']}
- Ref 1: {parsed_data['bill_info']['ref1']}
- Ref 2: {parsed_data['bill_info']['ref2']}

สรุป:
- จำนวนรายการ: {parsed_data['summary']['total_items']}
- Closed: {parsed_data['summary']['closed_items']} / Active: {parsed_data['summary']['active_items']}
- ยอดรวมจากเอกสาร: {parsed_data['summary']['total_amount']} บาท
- ยอดรวมคำนวณจากรายการ: {parsed_data['summary']['calculated_total']} บาท
- จำนวนหน้าทั้งหมด: {parsed_data['summary']['total_pdf_pages']}
- จำนวนหน้าที่ประมวลผล: {parsed_data['summary']['pages_processed']}

ตัวอย่าง 5 รายการแรก:
"""
    for i, item in enumerate(parsed_data["items"][:5]):
        summary_text += f"  {i+1}. [{item['status']}] {item['chassis_number']} | {item['period']} | {item['days']}วัน | {item['interest_rate']}% | เงินต้น {item['principal']} | ดอกเบี้ย {item['amount_due']}\n"

    if len(parsed_data["items"]) > 5:
        summary_text += f"  ... (อีก {len(parsed_data['items']) - 5} รายการ)\n"

    system_prompt = """คุณเป็นผู้ตรวจสอบเอกสารทางการเงิน กรุณาตรวจสอบข้อมูลสรุปจากเอกสาร Bill For Collection และตอบเป็น JSON:
{
  "verification": {
    "data_completeness": "เปอร์เซ็นต์ความสมบูรณ์",
    "amount_check": "ผลตรวจสอบยอดรวม",
    "issues_found": ["ปัญหาที่พบ"],
    "notes": "หมายเหตุ",
    "recommendation": "คำแนะนำ"
  }
}
ตอบเฉพาะ JSON เท่านั้น"""

    headers = {
        "Authorization": f"Bearer {TYPHOON_API_KEY}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": TYPHOON_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": summary_text},
        ],
        "max_tokens": 2048,
        "temperature": 0.1,
    }

    async with httpx.AsyncClient(timeout=120.0) as client:
        try:
            response = await client.post(TYPHOON_API_URL, headers=headers, json=payload)

            if response.status_code != 200:
                return _default_verification(f"Typhoon API error: {response.status_code}")

            result = response.json()
            content = result["choices"][0]["message"]["content"]

            # Parse JSON from response
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]

            return json.loads(content.strip())
        except Exception as e:
            return _default_verification(str(e))


def _default_verification(error_msg: str = "") -> dict:
    return {
        "verification": {
            "data_completeness": "N/A",
            "amount_check": "ไม่สามารถตรวจสอบได้",
            "issues_found": [error_msg] if error_msg else [],
            "notes": "ข้อมูลที่แสดงมาจากการแยกข้อมูลตารางโดยตรงจาก PDF",
            "recommendation": "กรุณาตรวจสอบข้อมูลด้วยตนเอง",
        }
    }


def _extract_text_from_pdf(pdf_path: str, max_pages: int = 0, password: str = "") -> tuple[str, int, int, list[str]]:
    """Extract plain text from PDF pages for generic Typhoon analysis."""
    open_kwargs = {}
    if password:
        open_kwargs["password"] = password

    pages_text = []
    total_pdf_pages = 0

    with pdfplumber.open(pdf_path, **open_kwargs) as pdf:
        total_pdf_pages = len(pdf.pages)
        pages_to_process = total_pdf_pages if max_pages <= 0 else min(max_pages, total_pdf_pages)
        for page_idx in range(pages_to_process):
            page = pdf.pages[page_idx]
            pages_text.append(page.extract_text() or "")

    raw_text = "\n\n--- หน้า ---\n\n".join(pages_text)
    return raw_text, total_pdf_pages, pages_to_process, pages_text


def _extract_text_with_progress(
    pdf_path: str,
    max_pages: int = 0,
    password: str = "",
    api_key: str = "",
    force_ocr: bool = False,
    progress_callback=None,
) -> tuple[str, int, int, list[str], list[dict]]:
    """Extract text with smart OCR: use native text when good, batch OCR only for garbled pages."""
    open_kwargs = {}
    if password:
        open_kwargs["password"] = password

    pages_text = []
    page_timings = []
    total_pdf_pages = 0
    pages_to_process = 0

    with pdfplumber.open(pdf_path, **open_kwargs) as pdf:
        total_pdf_pages = len(pdf.pages)
        pages_to_process = total_pdf_pages if max_pages <= 0 else min(max_pages, total_pdf_pages)

        # Pass 1: extract native text and identify pages needing OCR
        native_texts = []
        pages_needing_ocr = []

        for page_idx in range(pages_to_process):
            page = pdf.pages[page_idx]
            page_text = page.extract_text() or ""
            page_no = page_idx + 1

            should_ocr = force_ocr or _is_garbled_cid_text(page_text)
            if should_ocr:
                pages_needing_ocr.append((page_no, page_idx))
            native_texts.append(page_text)

        # Pass 2: batch OCR for pages that need it (1 API call instead of N)
        batch_elapsed = 0.0
        if pages_needing_ocr:
            page_nos = [p[0] for p in pages_needing_ocr]
            batch_start = time.perf_counter()
            ocr_texts = _call_typhoon_ocr_batch(
                pdf_path, api_key or TYPHOON_API_KEY, page_nos
            )
            batch_elapsed = round(time.perf_counter() - batch_start, 2)
            per_page_sec = round(batch_elapsed / max(len(page_nos), 1), 2)

            for idx, (page_no, page_idx) in enumerate(pages_needing_ocr):
                if idx < len(ocr_texts) and ocr_texts[idx]:
                    native_texts[page_idx] = ocr_texts[idx]

                if progress_callback:
                    progress_callback(page_no, pages_to_process, page_no, per_page_sec)

        # Build final output with timings
        ocr_page_set = {p[0] for p in pages_needing_ocr}
        for page_idx in range(pages_to_process):
            page_no = page_idx + 1
            page_text = _sanitize_text_for_display(native_texts[page_idx])
            pages_text.append(page_text)

            if page_no in ocr_page_set and pages_needing_ocr:
                elapsed = round(batch_elapsed / len(pages_needing_ocr), 2)
            else:
                elapsed = 0.1
            page_timings.append({"page_number": page_no, "elapsed_seconds": elapsed})

            if progress_callback and page_no not in ocr_page_set:
                progress_callback(page_no, pages_to_process, page_no, elapsed)

    raw_text = "\n\n--- หน้า ---\n\n".join(pages_text)
    return raw_text, total_pdf_pages, pages_to_process, pages_text, page_timings


def _init_tax_ocr_job(job_id: str):
    with TAX_OCR_JOBS_LOCK:
        TAX_OCR_JOBS[job_id] = {
            "status": "running",
            "message": "กำลังเตรียมไฟล์",
            "current_step": 0,
            "total_steps": 0,
            "current_page_number": 0,
            "page_timings": [],
            "result": None,
            "error": "",
        }


def _update_tax_ocr_job(job_id: str, **kwargs):
    with TAX_OCR_JOBS_LOCK:
        job = TAX_OCR_JOBS.get(job_id)
        if not job:
            return
        job.update(kwargs)


def _sanitize_text_for_display(text: str) -> str:
    """Clean noisy OCR/native text for UI display."""
    if not text:
        return ""
    cleaned = re.sub(r"\(cid:\d+\)", " ", text)
    cleaned = cleaned.replace("□", " ").replace("�", " ")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _is_garbled_cid_text(text: str) -> bool:
    if not text:
        return True
    cid_count = len(re.findall(r"\(cid:\d+\)", text))
    bad_glyph_count = text.count("□") + text.count("�")
    if cid_count >= 8:
        return True
    if bad_glyph_count >= 6:
        return True
    if cid_count > 0:
        ratio = cid_count / max(len(text.split()), 1)
        return ratio > 0.08
    if bad_glyph_count > 0:
        ratio = bad_glyph_count / max(len(text), 1)
        return ratio > 0.01
    return False


def _call_typhoon_ocr_page(file_path: str, api_key: str, page_number: int) -> str:
    """Call Typhoon OCR for a single page and return natural text."""
    results = _call_typhoon_ocr_batch(file_path, api_key, [page_number])
    return results[0] if results else ""


def _call_typhoon_ocr_batch(
    file_path: str, api_key: str, page_numbers: list[int]
) -> list[str]:
    """Call Typhoon OCR for multiple pages in a single request. Returns list of texts per page."""
    if not api_key or not page_numbers:
        return [""] * len(page_numbers)
    headers = {"Authorization": f"Bearer {api_key}"}
    data = {
        "model": TYPHOON_OCR_MODEL,
        "task_type": "default",
        "max_tokens": "16384",
        "temperature": "0.1",
        "top_p": "0.6",
        "repetition_penalty": "1.2",
        "pages": json.dumps(page_numbers),
    }

    with open(file_path, "rb") as file_obj:
        files = {"file": (os.path.basename(file_path), file_obj, "application/pdf")}
        with httpx.Client(timeout=180.0) as client:
            response = client.post(TYPHOON_OCR_URL, headers=headers, data=data, files=files)

    if response.status_code != 200:
        return [""] * len(page_numbers)

    try:
        result = response.json()
    except Exception:
        return [""] * len(page_numbers)

    texts = []
    for page_result in result.get("results", []):
        if page_result.get("success") and page_result.get("message"):
            content = page_result["message"]["choices"][0]["message"]["content"]
            try:
                parsed_content = json.loads(content)
                text = parsed_content.get("natural_text", content)
            except json.JSONDecodeError:
                text = content
            texts.append(text if text else "")
        else:
            texts.append("")
    return texts


def _normalize_for_lookup(value: str) -> str:
    return re.sub(r"[^0-9a-z]+", "", (value or "").lower())


def _find_token_after_label(
    lines: list[str],
    label_keywords: list[str],
    token_pattern: str,
    lookahead_lines: int = 2,
) -> str:
    normalized_keywords = [_normalize_for_lookup(keyword) for keyword in label_keywords]
    token_re = re.compile(token_pattern, flags=re.IGNORECASE)
    for index, line in enumerate(lines):
        normalized_line = _normalize_for_lookup(line)
        if not any(keyword in normalized_line for keyword in normalized_keywords):
            continue
        inline_match = token_re.search(line)
        if inline_match:
            return inline_match.group(1) if inline_match.lastindex else inline_match.group(0)
        for candidate in lines[index + 1: index + 1 + max(lookahead_lines, 0)]:
            candidate_match = token_re.search(candidate)
            if candidate_match:
                return candidate_match.group(1) if candidate_match.lastindex else candidate_match.group(0)
    return ""


def _find_value_after_label(lines: list[str], label_keywords: list[str]) -> str:
    money_pattern = re.compile(r"[0-9][0-9,]*\.[0-9]{2}\)?")
    normalized_keywords = [_normalize_for_lookup(keyword) for keyword in label_keywords]
    for index, line in enumerate(lines):
        normalized_line = _normalize_for_lookup(line)
        if not any(keyword in normalized_line for keyword in normalized_keywords):
            continue
        for candidate in lines[index + 1: index + 4]:
            match = money_pattern.search(candidate)
            if match:
                return match.group(0)
    return ""


def _extract_amounts_from_totals_block(lines: list[str]) -> dict[str, str]:
    money_pattern = re.compile(r"([0-9][0-9,]*\.[0-9]{2}\)?)")

    def detect_label(normalized_line: str) -> str:
        if "totalamount" in normalized_line:
            return "total_amount"
        if "discount" in normalized_line:
            return "discount"
        if "beforevat" in normalized_line:
            return "before_vat"
        if "withholdingtax" in normalized_line:
            return "withholding_tax"
        if "netamount" in normalized_line:
            return "net_amount"
        if "vat" in normalized_line and "beforevat" not in normalized_line:
            return "vat"
        if "nettotal" in normalized_line:
            return "net_total"
        return ""

    label_rows = []
    amounts = {}
    seen_labels = set()
    for index, line in enumerate(lines):
        label = detect_label(_normalize_for_lookup(line))
        if not label or label in seen_labels:
            continue
        seen_labels.add(label)
        label_rows.append((label, index))
        inline_match = money_pattern.search(line)
        if inline_match:
            amounts[label] = inline_match.group(1)

    if not label_rows:
        return amounts
    if any(amounts.get(key) for key in ["net_amount", "vat"]):
        return amounts

    last_label_index = label_rows[-1][1]
    value_candidates = []
    for line in lines[last_label_index + 1: last_label_index + 30]:
        match = money_pattern.search(line)
        if not match:
            continue
        value_candidates.append(match.group(1))
        if len(value_candidates) >= len(label_rows):
            break

    if len(value_candidates) >= len(label_rows):
        for idx, (label, _) in enumerate(label_rows):
            if not amounts.get(label):
                amounts[label] = value_candidates[idx]
    return amounts


def _extract_company_name(lines: list[str], text_blob: str) -> str:
    """Extract company/seller name from invoice text."""
    # Known OMODA patterns
    for line in lines:
        lower_line = line.lower()
        if "omoda" in lower_line and "jaecoo" in lower_line:
            return "บริษัท โอโมดา แอนด์ เจคู (ประเทศไทย) จำกัด (สำนักงานใหญ่)"

    # Generic: look for Thai company patterns (บริษัท/ห้างหุ้นส่วน ... จำกัด)
    company_pattern = re.compile(
        r"((?:บริษัท|ห้างหุ้นส่วน)\s*[^\n]{5,120}(?:จำกัด|สาขา|สำนักงานใหญ่)?(?:\s*\([^)]+\))?)",
        re.IGNORECASE,
    )
    match = company_pattern.search(text_blob)
    if match:
        return match.group(1).strip()

    # Look after "ชื่อผู้ขาย" / "Seller" / "ผู้ขาย" labels
    seller_keywords = ["ชื่อผู้ขาย", "seller", "ผู้ขาย", "ชื่อผู้จำหน่าย", "ผู้จำหน่าย"]
    for idx, line in enumerate(lines):
        line_str = line or ""
        line_lower = line_str.lower()
        if not any(kw in line_str or kw in line_lower for kw in seller_keywords):
            continue
        for cand in lines[idx + 1 : idx + 4]:
            c = cand.strip()
            if len(c) >= 6 and not re.match(r"^[0-9\s,\.\-/]+$", c):
                return c

    # Fallback: first substantial non-numeric line (typical company block at top)
    skip_pattern = re.compile(
        r"^(tax\s*invoice|ใบกำกับภาษี|bill|เลขที่|date|วันที่|tax\s*id|เลขประจำตัว|net\s*amount|vat|ภาษีมูลค่าเพิ่ม|[0-9]{7,}|[0-9,]+\.\d{2})\b",
        re.IGNORECASE,
    )
    for line in lines[:25]:
        cleaned = line.strip()
        if len(cleaned) < 6:
            continue
        if skip_pattern.match(cleaned):
            continue
        if re.match(r"^[0-9\s,\.\-/]+$", cleaned):
            continue
        if cleaned.startswith("(") and "cid:" in cleaned.lower():
            continue
        # Likely company line
        return cleaned

    return ""


def extract_tax_invoice_row(extracted_text: str) -> dict[str, str] | None:
    """Tax invoice parser adapted from OCR OMODA project."""
    lines = [line.strip() for line in (extracted_text or "").splitlines() if line.strip()]
    if not lines:
        return None
    text_blob = "\n".join(lines)

    company = _extract_company_name(lines, text_blob)

    tax_id = _find_token_after_label(
        lines,
        ["TaxID", "Tax ID", "เลขประจำตัวผู้เสียภาษี", "เลขที่ประจำตัวผู้เสียภาษี"],
        r"([0-9]{10,13})",
    )
    if not tax_id:
        tax_id_match = re.search(r"Tax\s*ID\s*:\s*([0-9]{10,13})", text_blob, flags=re.IGNORECASE)
        tax_id = tax_id_match.group(1) if tax_id_match else ""

    if not company and tax_id == "0105567128630":
        company = "บริษัท โอโมดา แอนด์ เจคู (ประเทศไทย) จำกัด (สำนักงานใหญ่)"

    invoice_no = ""
    invoice_match = re.search(
        r"Tax\s*Invoice[^\n\r0-9]{0,30}([0-9]{7,12})",
        text_blob,
        flags=re.IGNORECASE,
    )
    if invoice_match:
        invoice_no = invoice_match.group(1)
    if not invoice_no:
        invoice_line_indexes = [
            idx for idx, line in enumerate(lines) if "taxinvoice" in _normalize_for_lookup(line)
        ]
        for idx in invoice_line_indexes:
            window = " ".join(lines[idx: idx + 3])
            candidate = re.search(r"\b([0-9]{7,12})\b", window)
            if candidate:
                invoice_no = candidate.group(1)
                break

    date_value = _find_token_after_label(lines, ["Date", "วันที่"], r"(\d{2}[/\-]\d{2}[/\-]\d{4})")
    if not date_value:
        date_match = re.search(
            r"Date\s*:\s*(\d{2}[/\-]\d{2}[/\-]\d{4})",
            text_blob,
            flags=re.IGNORECASE,
        )
        date_value = date_match.group(1) if date_match else ""
    if not date_value:
        date_match = re.search(r"(\d{2}[/\-]\d{2}[/\-]\d{4})", text_blob)
        date_value = date_match.group(1) if date_match else ""

    totals_map = _extract_amounts_from_totals_block(lines)
    net_amount = totals_map.get("net_amount", "") or _find_value_after_label(
        lines, ["Net Amount", "ยอดเงินสุทธิ"]
    )
    vat_amount = totals_map.get("vat", "") or _find_value_after_label(
        lines, ["Vat", "ภาษีมูลค่าเพิ่ม"]
    )

    row = {
        "company": company,
        "tax_id": tax_id,
        "invoice_no": invoice_no,
        "date": date_value,
        "net_amount": net_amount,
        "vat_amount": vat_amount,
    }
    if not any([row["tax_id"], row["invoice_no"], row["date"]]):
        return None
    return row


def extract_tax_invoice_rows_from_page_texts(page_texts: list[str]) -> list[dict[str, str]]:
    rows = []
    seen = set()
    for page_text in page_texts or []:
        row = extract_tax_invoice_row(page_text)
        if not row:
            continue
        key = (
            row.get("company", ""),
            row.get("tax_id", ""),
            row.get("invoice_no", ""),
            row.get("date", ""),
            row.get("net_amount", ""),
            row.get("vat_amount", ""),
        )
        if key in seen:
            continue
        seen.add(key)
        rows.append(row)
    return rows


def _tax_analysis_from_row(row: dict[str, str], total_pages: int, pages_processed: int) -> dict:
    """Build UI-compatible analysis object from parsed tax row."""
    amount_text = row.get("net_amount", "") or "0.00"
    return {
        "bill_info": {
            "bill_number": row.get("invoice_no", ""),
            "customer_name": row.get("company", ""),
            "date": row.get("date", ""),
            "address": "",
            "credit_line_no": row.get("tax_id", ""),
            "ref1": "",
            "ref2": "",
        },
        "items": [],
        "summary": {
            "total_items": 1,
            "total_amount": amount_text,
            "total_amount_text": "",
            "closed_items": 0,
            "active_items": 1,
            "bank_name": "",
            "calculated_total": amount_text,
            "total_pdf_pages": total_pages,
            "pages_processed": pages_processed,
        },
        "verification": _default_verification("สรุปจาก OCR parser (OMODA)")["verification"],
    }


def _extract_json_from_model_response(content: str) -> dict:
    """Extract JSON object from model response content."""
    if "```json" in content:
        content = content.split("```json", 1)[1].split("```", 1)[0]
    elif "```" in content:
        content = content.split("```", 1)[1].split("```", 1)[0]
    return json.loads(content.strip())


async def analyze_tax_with_typhoon(
    raw_text: str,
    filename: str,
    total_pages: int,
    pages_processed: int,
    api_key_override: str = "",
) -> dict:
    """Analyze TAX INVOICE/RECEIPT text and return UI-compatible structure."""
    safe_text = raw_text[:10000] if raw_text else ""

    used_api_key = api_key_override.strip() if api_key_override else TYPHOON_API_KEY

    # If Typhoon key is missing, return a safe fallback structure so UI can continue.
    if not used_api_key or used_api_key == "your_typhoon_api_key_here":
        return {
            "bill_info": {
                "bill_number": "",
                "customer_name": "",
                "date": "",
                "address": "",
                "credit_line_no": "",
                "ref1": "",
                "ref2": "",
            },
            "items": [],
            "summary": {
                "total_items": 0,
                "total_amount": "0.00",
                "total_amount_text": "",
                "closed_items": 0,
                "active_items": 0,
                "bank_name": "",
                "calculated_total": "0.00",
                "total_pdf_pages": total_pages,
                "pages_processed": pages_processed,
            },
            "verification": _default_verification("Typhoon API key not configured")["verification"],
        }

    system_prompt = """คุณคือผู้ช่วยวิเคราะห์เอกสาร TAX INVOICE/RECEIPT
กรุณาอ่านข้อความจากเอกสารแล้วสรุปผลออกมาเป็น JSON ตาม schema นี้เท่านั้น:
{
  "bill_info": {
    "bill_number": "เลขที่เอกสาร",
    "customer_name": "ชื่อลูกค้า/ผู้ซื้อ",
    "date": "วันที่เอกสาร",
    "address": "ที่อยู่ลูกค้า",
    "credit_line_no": "เลขวงเงิน/เลขอ้างอิงถ้ามี",
    "ref1": "อ้างอิง 1",
    "ref2": "อ้างอิง 2"
  },
  "summary": {
    "total_items": 0,
    "total_amount": "0.00",
    "total_amount_text": "",
    "closed_items": 0,
    "active_items": 0,
    "bank_name": "",
    "calculated_total": "0.00"
  },
  "items": [
    {
      "chassis_number": "",
      "status": "Active",
      "description": "",
      "period": "",
      "days": "",
      "interest_rate": "",
      "due_date": "",
      "principal": "",
      "amount_due": "0.00",
      "page_number": 1
    }
  ],
  "verification": {
    "data_completeness": "0%",
    "amount_check": "",
    "issues_found": [],
    "notes": "",
    "recommendation": ""
  }
}
กติกา:
- ตอบเฉพาะ JSON ห้ามมีคำอธิบายเพิ่ม
- ถ้าไม่มีข้อมูลในฟิลด์ให้ใส่ค่าว่างหรือ 0 ตามชนิดข้อมูล
- total_amount และ calculated_total ให้คงรูปแบบสตริงตัวเลข เช่น "219,950.62"
"""

    user_prompt = f"""ไฟล์: {filename}
จำนวนหน้าทั้งหมด: {total_pages}
จำนวนหน้าที่ประมวลผล: {pages_processed}

ข้อความที่อ่านได้จากเอกสาร:
{safe_text}
"""

    headers = {
        "Authorization": f"Bearer {used_api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": TYPHOON_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "max_tokens": 2048,
        "temperature": 0.1,
    }

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(TYPHOON_API_URL, headers=headers, json=payload)
        if response.status_code != 200:
            raise ValueError(f"Typhoon API error: {response.status_code}")

        content = response.json()["choices"][0]["message"]["content"]
        model_data = _extract_json_from_model_response(content)

        summary = model_data.get("summary", {})
        summary["total_pdf_pages"] = total_pages
        summary["pages_processed"] = pages_processed

        return {
            "bill_info": model_data.get("bill_info", {}),
            "items": model_data.get("items", []),
            "summary": summary,
            "verification": model_data.get("verification", _default_verification()["verification"]),
        }
    except Exception as e:
        return {
            "bill_info": {
                "bill_number": "",
                "customer_name": "",
                "date": "",
                "address": "",
                "credit_line_no": "",
                "ref1": "",
                "ref2": "",
            },
            "items": [],
            "summary": {
                "total_items": 0,
                "total_amount": "0.00",
                "total_amount_text": "",
                "closed_items": 0,
                "active_items": 0,
                "bank_name": "",
                "calculated_total": "0.00",
                "total_pdf_pages": total_pages,
                "pages_processed": pages_processed,
            },
            "verification": _default_verification(str(e))["verification"],
        }


# ======== ROUTES ========

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/api/check-pdf")
async def check_pdf(file: UploadFile = File(...)):
    """Check if a PDF file is encrypted/password-protected."""
    if not file.filename.lower().endswith(".pdf"):
        return JSONResponse(content={"needs_password": False, "error": "ไม่ใช่ไฟล์ PDF"})

    temp_path = upload_dir / f"check_{file.filename}"
    try:
        contents = await file.read()
        with open(temp_path, "wb") as f:
            f.write(contents)

        try:
            # Try to open without password
            with pdfplumber.open(str(temp_path)) as pdf:
                # Accessing pages usually triggers decryption attempt
                if len(pdf.pages) == 0:
                     return JSONResponse(content={"needs_password": True})
                
                # Try to extract text from first page - some encrypted files open but fail on read
                try:
                    pdf.pages[0].extract_text()
                except Exception:
                     # Failed to read content
                     return JSONResponse(content={"needs_password": True})
            
            return JSONResponse(content={"needs_password": False})
            
        except Exception as e:
            error_msg = str(e).lower()
            print(f"PDF Check Error for {file.filename}: {error_msg}")
            
            # Assume any error during open means it might be encrypted or corrupted
            # It's safer to ask for password than to fail silently
            return JSONResponse(content={"needs_password": True})

    except Exception as e:
         return JSONResponse(content={"needs_password": False, "error": f"Upload error: {str(e)}"})
    finally:
        if temp_path.exists():
            try:
                temp_path.unlink()
            except:
                pass


@app.post("/api/scan")
async def scan_document(
    file: UploadFile = File(...),
    max_pages: int = Form(default=0),
    pdf_password: str = Form(default=""),
):
    """Upload and scan a PDF document.
    max_pages=0 means process all pages.
    pdf_password: password for encrypted PDF files.
    """
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="กรุณาอัปโหลดไฟล์ PDF เท่านั้น")

    temp_path = upload_dir / file.filename
    try:
        contents = await file.read()
        with open(temp_path, "wb") as f:
            f.write(contents)

        # Step 1: Extract structured data from PDF tables
        try:
            parsed = extract_from_pdf(str(temp_path), max_pages=max_pages, password=pdf_password)
        except Exception as e:
            error_msg = str(e).lower()
            if "password" in error_msg or "encrypt" in error_msg or "decrypt" in error_msg:
                raise HTTPException(
                    status_code=401,
                    detail="ไฟล์ PDF มีรหัสผ่าน กรุณาใส่รหัสผ่านที่ถูกต้อง",
                )
            raise

        if not parsed["items"]:
            raise HTTPException(
                status_code=400,
                detail="ไม่พบข้อมูลตารางในไฟล์ PDF",
            )

        # Step 2: Verify with Typhoon AI
        verification = _default_verification()
        if TYPHOON_API_KEY and TYPHOON_API_KEY != "your_typhoon_api_key_here":
            verification = await verify_with_typhoon(parsed)

        # Build response
        analysis = {
            "bill_info": parsed["bill_info"],
            "items": parsed["items"],
            "summary": parsed["summary"],
            "verification": verification.get("verification", {}),
        }

        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "extracted_text_preview": parsed["raw_text"][:2000],
            "total_pages": parsed["summary"]["total_pdf_pages"],
            "pages_processed": parsed["summary"]["pages_processed"],
            "analysis": analysis,
        })

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"เกิดข้อผิดพลาด: {str(e)}")
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _run_tax_ocr_job(
    job_id: str,
    uploads: list[tuple[Path, str]],
    max_pages: int,
    pdf_password: str,
    api_key: str,
):
    try:
        total_files = len(uploads)
        combined_page_texts = []
        combined_page_timings = []
        combined_texts = []
        total_elapsed = 0.0
        page_number_offset = 0

        for file_index, (temp_path, filename) in enumerate(uploads, start=1):
            def on_progress(step: int, total: int, page_number: int, elapsed: float):
                with TAX_OCR_JOBS_LOCK:
                    job = TAX_OCR_JOBS.get(job_id)
                    if not job:
                        return
                    global_page = page_number_offset + page_number
                    timings = job.get("page_timings", [])
                    timings.append({"page_number": global_page, "elapsed_seconds": elapsed})
                    job["page_timings"] = timings
                    safe_total = max(total, 1)
                    job["status"] = "running"
                    job["message"] = f"ไฟล์ {file_index}/{total_files} กำลัง OCR หน้า {step}/{safe_total}"
                    job["current_step"] = ((file_index - 1) * 100) + int((step / safe_total) * 100)
                    job["total_steps"] = max(total_files * 100, 1)
                    job["current_page_number"] = global_page

            started = time.perf_counter()
            raw_text, total_pdf_pages, pages_processed, page_texts, _ = _extract_text_with_progress(
                str(temp_path),
                max_pages=max_pages,
                password=pdf_password,
                api_key=api_key or TYPHOON_API_KEY,
                force_ocr=False,  # Use native text when good; OCR only garbled pages (faster)
                progress_callback=on_progress,
            )
            elapsed = round(time.perf_counter() - started, 2)
            total_elapsed += elapsed

            combined_texts.append(raw_text)
            combined_page_texts.extend(page_texts)
            for idx in range(len(page_texts)):
                combined_page_timings.append({
                    "page_number": page_number_offset + idx + 1,
                    "elapsed_seconds": round(elapsed / max(len(page_texts), 1), 2),
                })
            page_number_offset += len(page_texts)
            # keep compatibility if max_pages limits pages
            _ = (total_pdf_pages, pages_processed)

        merged_text = "\n\n".join([txt for txt in combined_texts if txt])
        total_pages = len(combined_page_texts)
        analysis = asyncio.run(
            analyze_tax_with_typhoon(
                raw_text=merged_text,
                filename=", ".join([name for _, name in uploads]),
                total_pages=total_pages,
                pages_processed=total_pages,
                api_key_override=api_key,
            )
        )

        tax_rows = extract_tax_invoice_rows_from_page_texts(combined_page_texts)
        bill_info = analysis.get("bill_info") or {}
        bill_date = bill_info.get("date", "").strip()
        bill_number = (bill_info.get("bill_number") or "").strip()
        for row in tax_rows:
            if not row.get("date") and bill_date and bill_number:
                inv = (row.get("invoice_no") or "").strip()
                if inv == bill_number:
                    row["date"] = bill_date
                    break
        if tax_rows:
            first = tax_rows[0]
            parser_analysis = _tax_analysis_from_row(first, total_pages, total_pages)
            merged_bill = analysis.get("bill_info", {}) or {}
            merged_bill["bill_number"] = merged_bill.get("bill_number") or parser_analysis["bill_info"]["bill_number"]
            merged_bill["customer_name"] = merged_bill.get("customer_name") or parser_analysis["bill_info"]["customer_name"]
            merged_bill["date"] = merged_bill.get("date") or parser_analysis["bill_info"]["date"]
            merged_bill["credit_line_no"] = merged_bill.get("credit_line_no") or parser_analysis["bill_info"]["credit_line_no"]
            analysis["bill_info"] = merged_bill

        result = {
            "success": True,
            "filename": ", ".join([name for _, name in uploads]),
            "extracted_text_preview": merged_text[:2000],
            "page_texts": combined_page_texts,
            "page_htmls": combined_page_texts,
            "total_pages": total_pages,
            "pages_processed": total_pages,
            "analysis": analysis,
            "tax_invoice_rows": tax_rows,
            "page_timings": combined_page_timings,
            "elapsed_seconds": round(total_elapsed, 2),
        }
        _update_tax_ocr_job(
            job_id,
            status="completed",
            message="OCR เสร็จแล้ว",
            result=result,
            current_step=max(total_files * 100, 1),
            total_steps=max(total_files * 100, 1),
            current_page_number=total_pages if total_pages > 0 else 0,
        )
    except Exception as e:
        _update_tax_ocr_job(job_id, status="failed", message="OCR ล้มเหลว", error=str(e))
    finally:
        for temp_path, _ in uploads:
            if temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception:
                    pass


@app.post("/api/tax-ocr/start")
async def tax_ocr_start(
    pdf_file: list[UploadFile] = File(default=[]),
    file: Optional[UploadFile] = File(default=None),
    max_pages: int = Form(default=0),
    pdf_password: str = Form(default=""),
    api_key: str = Form(default=""),
):
    uploaded = [f for f in (pdf_file or []) if f and f.filename]
    if file and file.filename:
        uploaded.append(file)
    if not uploaded:
        raise HTTPException(status_code=400, detail="กรุณาเลือกไฟล์ PDF ก่อน")
    for f in uploaded:
        if not f.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="กรุณาอัปโหลดไฟล์ PDF เท่านั้น")

    job_id = f"tax_{int(time.time() * 1000)}_{os.getpid()}"
    _init_tax_ocr_job(job_id)
    uploads = []
    for idx, up in enumerate(uploaded, start=1):
        temp_path = upload_dir / f"{job_id}_{idx}_{up.filename}"
        contents = await up.read()
        with open(temp_path, "wb") as f:
            f.write(contents)
        uploads.append((temp_path, up.filename))

    worker = threading.Thread(
        target=_run_tax_ocr_job,
        args=(job_id, uploads, max_pages, pdf_password, api_key),
        daemon=True,
    )
    worker.start()
    return {"ok": True, "job_id": job_id}


@app.get("/api/tax-ocr/status/{job_id}")
async def tax_ocr_status(job_id: str):
    with TAX_OCR_JOBS_LOCK:
        job = TAX_OCR_JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="ไม่พบงาน OCR")
        return {
            "ok": True,
            "status": job.get("status", "running"),
            "message": job.get("message", ""),
            "current_step": job.get("current_step", 0),
            "total_steps": job.get("total_steps", 0),
            "current_page_number": job.get("current_page_number", 0),
            "page_timings": job.get("page_timings", []),
            "error": job.get("error", ""),
            "result": job.get("result") if job.get("status") == "completed" else None,
        }


@app.post("/api/scan-tax")
async def scan_tax_document(
    file: UploadFile = File(...),
    max_pages: int = Form(default=0),
    pdf_password: str = Form(default=""),
):
    """Upload and scan TAX INVOICE/RECEIPT PDF document via Typhoon."""
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="กรุณาอัปโหลดไฟล์ PDF เท่านั้น")

    temp_path = upload_dir / f"tax_{file.filename}"
    try:
        contents = await file.read()
        with open(temp_path, "wb") as f:
            f.write(contents)

        # OCR-style text extraction (aligned with OCR OMODA flow)
        try:
            raw_text, total_pdf_pages, pages_processed, page_texts = _extract_text_from_pdf(
                str(temp_path), max_pages=max_pages, password=pdf_password
            )
        except Exception as e:
            error_msg = str(e).lower()
            if "password" in error_msg or "encrypt" in error_msg or "decrypt" in error_msg:
                raise HTTPException(
                    status_code=401,
                    detail="ไฟล์ PDF มีรหัสผ่าน กรุณาใส่รหัสผ่านที่ถูกต้อง",
                )
            raise

        # First pass with tax parser from OCR OMODA.
        parser_analysis = None
        tax_row = extract_tax_invoice_row(raw_text)
        if tax_row:
            parser_analysis = _tax_analysis_from_row(tax_row, total_pdf_pages, pages_processed)

        # Then send to Typhoon for richer verification and structure.
        analysis = await analyze_tax_with_typhoon(
            raw_text=raw_text,
            filename=file.filename,
            total_pages=total_pdf_pages,
            pages_processed=pages_processed,
        )
        if parser_analysis:
            # Keep Typhoon result as primary but fill critical invoice fields from parser when missing.
            merged_bill = analysis.get("bill_info", {}) or {}
            merged_bill["bill_number"] = merged_bill.get("bill_number") or parser_analysis["bill_info"]["bill_number"]
            merged_bill["customer_name"] = merged_bill.get("customer_name") or parser_analysis["bill_info"]["customer_name"]
            merged_bill["date"] = merged_bill.get("date") or parser_analysis["bill_info"]["date"]
            merged_bill["credit_line_no"] = merged_bill.get("credit_line_no") or parser_analysis["bill_info"]["credit_line_no"]
            analysis["bill_info"] = merged_bill

            merged_summary = analysis.get("summary", {}) or {}
            if not merged_summary.get("total_amount") or str(merged_summary.get("total_amount")).strip() in {"", "0", "0.0", "0.00"}:
                merged_summary["total_amount"] = parser_analysis["summary"]["total_amount"]
                merged_summary["calculated_total"] = parser_analysis["summary"]["calculated_total"]
            merged_summary["total_pdf_pages"] = total_pdf_pages
            merged_summary["pages_processed"] = pages_processed
            analysis["summary"] = merged_summary

        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "extracted_text_preview": raw_text[:2000],
            "page_texts": page_texts,
            "total_pages": total_pdf_pages,
            "pages_processed": pages_processed,
            "analysis": analysis,
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"เกิดข้อผิดพลาด: {str(e)}")
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _export_tax_to_docx(tax_rows: list[dict]) -> io.BytesIO:
    """Export tax_invoice_rows to DOCX document."""
    headers = ["บริษัท", "เลขที่ประจำตัวผู้เสียภาษี TaxID", "เลขที่ใบกำกับภาษี Tax Invoice", "วันที่ Date", "ยอดเงินสุทธิ / Net Amount", "ภาษีมูลค่าเพิ่ม / Vat"]
    cols = ["company", "tax_id", "invoice_no", "date", "net_amount", "vat_amount"]
    doc = Document()
    doc.add_heading("Tax Invoice OCR Result", level=1)
    if tax_rows:
        table = doc.add_table(rows=len(tax_rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        for c_idx, h in enumerate(headers):
            table.cell(0, c_idx).text = h
            for run in table.cell(0, c_idx).paragraphs[0].runs:
                run.bold = True
        for r_idx, row in enumerate(tax_rows, start=1):
            for c_idx, col in enumerate(cols):
                table.cell(r_idx, c_idx).text = str(row.get(col, ""))
    else:
        doc.add_paragraph("ยังไม่มีข้อมูล")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _get_tax_invoice_connection_params() -> dict:
    """Parse DATABASE_URL to get server and driver for TaxInvoiceDB."""
    url = os.getenv("DATABASE_URL", "")
    if not url:
        return {"server": "LAPTOP-2CN8L0R4\\SQLEXPRESS", "driver": "ODBC Driver 17 for SQL Server", "trusted": True}
    # mssql+pyodbc://SERVER/DATABASE?driver=ODBC+Driver+17+for+SQL+Server&trusted_connection=yes
    server = "LAPTOP-2CN8L0R4\\SQLEXPRESS"
    driver = "ODBC Driver 17 for SQL Server"
    trusted = True
    ms = re.search(r"mssql\+pyodbc://([^/]+)/", url)
    if ms:
        server = ms.group(1).strip()
    md = re.search(r"[?&]driver=([^&]+)", url)
    if md:
        driver = md.group(1).replace("+", " ").strip()
    mt = re.search(r"[?&]trusted_connection=([^&]+)", url, re.I)
    if mt:
        trusted = str(mt.group(1)).lower() in ("yes", "true", "1")
    return {"server": server, "driver": driver, "trusted": trusted}


def _ensure_tax_invoice_db(conn_params: dict) -> None:
    """Create TaxInvoiceDB database if not exists."""
    conn_str = (
        f"Driver={{{conn_params['driver']}}};"
        f"Server={conn_params['server']};"
        f"Database=master;"
    )
    if conn_params.get("trusted", True):
        conn_str += "Trusted_Connection=yes;"
    conn = pyodbc.connect(conn_str, autocommit=True)
    try:
        cur = conn.cursor()
        cur.execute("SELECT name FROM sys.databases WHERE name = N'TaxInvoiceDB'")
        if not cur.fetchone():
            cur.execute("CREATE DATABASE TaxInvoiceDB")
    finally:
        conn.close()


def _ensure_tax_invoice_table(conn_params: dict) -> None:
    """Create TaxInvoice table if not exists."""
    conn_str = (
        f"Driver={{{conn_params['driver']}}};"
        f"Server={conn_params['server']};"
        f"Database=TaxInvoiceDB;"
    )
    if conn_params.get("trusted", True):
        conn_str += "Trusted_Connection=yes;"
    conn = pyodbc.connect(conn_str, autocommit=True)
    try:
        cur = conn.cursor()
        cur.execute("""
            IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = N'TaxInvoice')
            CREATE TABLE TaxInvoice (
                Id INT IDENTITY(1,1) PRIMARY KEY,
                [ชื่อไฟล์] NVARCHAR(500),
                [ลำดับ] INT,
                [บริษัท] NVARCHAR(500),
                [เลขที่ประจำตัวผู้เสียภาษี TaxID] NVARCHAR(50),
                [เลขที่ใบกำกับภาษี Tax Invoice] NVARCHAR(50),
                [วันที่ Date] NVARCHAR(50),
                [ยอดเงินสุทธิ / Net Amount] NVARCHAR(100),
                [ภาษีมูลค่าเพิ่ม / Vat] NVARCHAR(100)
            )
        """)
    finally:
        conn.close()


@app.post("/api/tax-import-database")
async def tax_import_database(request: Request):
    """Create TaxInvoiceDB database, table, and import tax invoice data."""
    try:
        data = await request.json()
        tax_rows = data.get("tax_invoice_rows") or []
        filename = data.get("filename") or ""

        if not tax_rows:
            raise HTTPException(status_code=400, detail="ไม่มีข้อมูลสำหรับ Import")

        params = _get_tax_invoice_connection_params()
        _ensure_tax_invoice_db(params)
        _ensure_tax_invoice_table(params)

        conn_str = (
            f"Driver={{{params['driver']}}};"
            f"Server={params['server']};"
            f"Database=TaxInvoiceDB;"
        )
        if params.get("trusted", True):
            conn_str += "Trusted_Connection=yes;"
        conn = pyodbc.connect(conn_str)
        try:
            cur = conn.cursor()
            for idx, row in enumerate(tax_rows, start=1):
                cur.execute("""
                    INSERT INTO TaxInvoice (
                        [ชื่อไฟล์], [ลำดับ], [บริษัท],
                        [เลขที่ประจำตัวผู้เสียภาษี TaxID],
                        [เลขที่ใบกำกับภาษี Tax Invoice],
                        [วันที่ Date],
                        [ยอดเงินสุทธิ / Net Amount],
                        [ภาษีมูลค่าเพิ่ม / Vat]
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    filename,
                    idx,
                    row.get("company") or "",
                    row.get("tax_id") or "",
                    row.get("invoice_no") or "",
                    row.get("date") or "",
                    row.get("net_amount") or "",
                    row.get("vat_amount") or "",
                ))
            conn.commit()
            return {"success": True, "rows_imported": len(tax_rows), "message": "Import ข้อมูลลง TaxInvoiceDB เรียบร้อยแล้ว"}
        finally:
            conn.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Import Database Error: {str(e)}")


@app.post("/api/tax-export/docx")
async def tax_export_docx(request: Request):
    """Export tax invoice data to DOCX (Word) file."""
    try:
        data = await request.json()
        tax_rows = data.get("tax_invoice_rows") or []
        if not tax_rows:
            raise HTTPException(status_code=400, detail="ไม่มีข้อมูลสำหรับ Export")
        docx_bytes = _export_tax_to_docx(tax_rows)
        filename = f"TAX_INVOICE_{datetime.date.today().isoformat()}.docx"
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export Error: {str(e)}")


@app.post("/api/import-db")
async def import_to_db(request: Request, db: Session = Depends(get_db)):
    """Save scanned analysis data to SQL Server."""
    def clean_num(val):
        if not val: return 0.0
        try:
            return float(str(val).replace(",", "").strip())
        except:
            return 0.0

    try:
        data = await request.json()
        analysis = data.get("analysis")
        if not analysis:
            raise HTTPException(status_code=400, detail="ไม่พบข้อมูล analysis")

        bill_info = analysis.get("bill_info", {})
        summary = analysis.get("summary", {})
        verification_data = analysis.get("verification", {})

        # 1. Create Bill record
        db_bill = Bill(
            BillNumber=bill_info.get("bill_number"),
            CustomerName=bill_info.get("customer_name"),
            BillDate=bill_info.get("date"),
            Address=bill_info.get("address"),
            CreditLineNo=bill_info.get("credit_line_no"),
            Ref1=bill_info.get("ref1"),
            Ref2=bill_info.get("ref2"),
            TotalAmount=clean_num(summary.get("total_amount")),
            CalculatedTotal=clean_num(summary.get("calculated_total")),
            BankName=summary.get("bank_name"),
            TotalPages=summary.get("total_pdf_pages")
        )
        db.add(db_bill)
        db.flush()  # To get db_bill.Id

        # 2. Create BillItems
        for item in analysis.get("items", []):
            db_item = BillItem(
                BillId=db_bill.Id,
                ChassisNumber=item.get("chassis_number"),
                Status=item.get("status"),
                Description=item.get("description", "ดอกเบี้ย"),
                Period=item.get("period"),
                Days=item.get("days"),
                InterestRate=item.get("interest_rate"),
                DueDate=item.get("due_date"),
                Principal=str(item.get("principal", "")),
                AmountDue=clean_num(item.get("amount_due")),
                PageNumber=item.get("page_number")
            )
            db.add(db_item)

        # 3. Create Verification record
        if verification_data:
            issues = verification_data.get("issues_found", [])
            issues_str = ", ".join(issues) if isinstance(issues, list) else str(issues)
            
            db_verify = Verification(
                BillId=db_bill.Id,
                DataCompleteness=verification_data.get("data_completeness"),
                AmountCheck=verification_data.get("amount_check"),
                Notes=f"{issues_str}\n\n{verification_data.get('notes', '')}".strip(),
                Recommendation=verification_data.get("recommendation")
            )
            db.add(db_verify)

        db.commit()
        return {"success": True, "bill_id": db_bill.Id}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Database Import Error: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
